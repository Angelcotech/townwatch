"""
Shared document-format detection + plain-text conversion.

Civic platforms serve meeting documents as PDF, DOCX, or legacy DOC depending
on jurisdiction vintage (CivicEngage AgendaCenter in particular serves DOCX
behind PDF-looking URLs). Every document extractor (agendas, minutes,
campaign finance, …) dispatches on the same content-type + magic-byte sniff,
so the detection and the format→text converters live here once.

Extractors keep their own `extract_from_document` entry points — the schema,
prompts, and recovery ladder differ per domain — but they all dispatch
through these helpers.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_CT = "application/msword"
PDF_CT = "application/pdf"


def normalize_ct(content_type: str | None) -> str:
    return (content_type or "").lower().split(";", 1)[0].strip()


def sniff_ct_from_magic(path: Path) -> str | None:
    """Sniff content type from the first bytes. Used when HTTP didn't tell us."""
    head = path.read_bytes()[:8]
    if head.startswith(b"%PDF"):
        return PDF_CT
    if head.startswith(b"PK\x03\x04"):  # ZIP container — likely DOCX
        return DOCX_CT
    if head.startswith(b"\xd0\xcf\x11\xe0"):  # OLE Compound Document — DOC, XLS, PPT
        return DOC_CT
    return None


def resolve_ct(path: Path, content_type: str | None) -> str:
    """The effective content type: trust a recognised HTTP header, otherwise
    sniff magic bytes. Returns "" when neither identifies the file."""
    ct = normalize_ct(content_type)
    if ct not in (PDF_CT, DOCX_CT, DOC_CT):
        ct = sniff_ct_from_magic(path) or ""
    return ct


def docx_to_text(doc_path: Path) -> str:
    """Read DOCX paragraphs + tables as plain text.

    DOCX is structured XML — no vision needed. Tables matter because many
    older meeting templates put the docket (or the motion/vote record) in a
    table rather than paragraph form.
    """
    from docx import Document  # lazy import — python-docx
    doc = Document(doc_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((c.text or "").strip() for c in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def doc_to_text(doc_path: Path) -> str:
    """Shell to libreoffice headless to convert legacy DOC → text.

    DOC is the binary Microsoft Word format; no good native Python reader
    exists. libreoffice is the standard converter on both Mac
    (brew install --cask libreoffice) and Linux (apt install libreoffice).
    """
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if binary is None:
        raise RuntimeError(
            "DOC extraction requires libreoffice. "
            "Install on Mac: brew install --cask libreoffice. "
            "Install on Linux: apt-get install libreoffice-core libreoffice-writer."
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run(
                [binary, "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(doc_path)],
                check=True, timeout=180, capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"libreoffice conversion failed: {e.stderr.decode('utf-8', errors='replace')[:500]}"
            ) from e
        txt_path = Path(tmpdir) / (doc_path.stem + ".txt")
        if not txt_path.exists():
            raise RuntimeError(f"libreoffice did not produce {txt_path.name}")
        return txt_path.read_text(encoding="utf-8", errors="replace")


def unsupported_format_error(path: Path, content_type: str | None) -> RuntimeError:
    return RuntimeError(
        f"unsupported document type {content_type!r} for {path.name} "
        f"(magic={path.read_bytes()[:4]!r})"
    )
