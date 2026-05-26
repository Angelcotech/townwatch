"""
Extract structured Campaign Contribution Disclosure Reports (CCDRs).

Designed to receive paper-filed records that arrive via Open Records Act
responses (Phase 1, pre-2027 for Georgia municipals) or, later, digital
exports from state ethics portals (e.g. ethics.ga.gov from 2027 onward).

Reuses the same content-type dispatch (PDF / DOCX / DOC) that the agenda
extractor uses. Schema is per-filing: one filer block + period block +
declared totals + list of contributions + list of expenditures.

**Status**: schema and prompts are complete; the API call is wired and
will work end-to-end. NOT YET validated against real CCDR documents.
The first 3-5 extractions should be reviewed manually before bulk
ingestion is trusted. See docstring on
`jobs/extract_campaign_finance.py` for the rollout protocol.
"""

from __future__ import annotations

import base64
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Literal

import anthropic
from pydantic import BaseModel, Field

from ..config import ANTHROPIC_API_KEY


# =====================================================================
# Pydantic schema — what Claude is constrained to return
# =====================================================================

FilingType = Literal[
    "pre_election",
    "post_election",
    "semi_annual",
    "two_business_day_large",
    "amendment",
    "final",
    "other",
]
ContributorType = Literal[
    "individual",
    "corporation",
    "pac",
    "union",
    "party",
    "candidate_committee",
    "self_funded",
    "other",
]
Confidence = Literal["high", "medium", "low"]


class CampaignFiler(BaseModel):
    candidate_name: str
    office_sought: str = Field(description="e.g. 'Mayor', 'City Council Member 2', 'Member, Board of Education District 3'")
    committee_name: str | None = Field(default=None, description="Campaign committee name if distinct from candidate")
    committee_address: str | None = None
    committee_treasurer: str | None = None


class FilingPeriod(BaseModel):
    election_cycle_year: int = Field(description="Year of the election cycle this filing covers, e.g. 2024")
    filing_type: FilingType
    period_start: str | None = Field(default=None, description="YYYY-MM-DD; start of the period the filing covers")
    period_end: str | None = Field(default=None, description="YYYY-MM-DD; end of the period the filing covers")
    filing_date: str | None = Field(default=None, description="YYYY-MM-DD; when the candidate actually filed the report")


class DeclaredTotals(BaseModel):
    total_contributions: float | None = Field(default=None, description="Total contributions declared on the filing's summary page for this period")
    total_expenditures: float | None = Field(default=None, description="Total expenditures declared for this period")
    cash_on_hand_start: float | None = None
    cash_on_hand_end: float | None = None
    in_kind_total: float | None = Field(default=None, description="Total in-kind contributions for this period")


class Contribution(BaseModel):
    contributor_name: str = Field(description="Exact name as written on the filing")
    contributor_type: ContributorType
    amount: float = Field(description="Cash + in-kind value of the contribution, in dollars")
    contribution_date: str | None = Field(default=None, description="YYYY-MM-DD when the contribution was made")
    contributor_address: str | None = None
    contributor_city: str | None = None
    contributor_state: str | None = Field(default=None, description="Two-letter USPS code")
    contributor_zip: str | None = None
    contributor_occupation: str | None = None
    contributor_employer: str | None = None
    is_in_kind: bool = Field(default=False, description="True if this is an in-kind contribution rather than cash")
    in_kind_description: str | None = None


class Expenditure(BaseModel):
    payee_name: str
    purpose: str | None = Field(default=None, description="Stated purpose of the expenditure")
    amount: float
    expenditure_date: str | None = Field(default=None, description="YYYY-MM-DD")
    payee_address: str | None = None


class CampaignFilingExtraction(BaseModel):
    filer: CampaignFiler
    period: FilingPeriod
    declared_totals: DeclaredTotals
    contributions: list[Contribution] = Field(default_factory=list)
    expenditures: list[Expenditure] = Field(default_factory=list)
    extraction_confidence: Confidence
    extraction_notes: str = Field(
        default="",
        description="Ambiguities, illegible passages, totals that don't reconcile, etc.",
    )


# =====================================================================
# Prompts
# =====================================================================

_RULES = """\
WHAT TO EXTRACT

This document is a Campaign Contribution Disclosure Report (CCDR) — a
periodic filing made by a candidate for public office to disclose
contributions received and expenditures made during a defined reporting
period. Extract:

  - filer: candidate name, office sought, committee name and treasurer
  - period: election cycle year, filing type (pre-election, post-election,
    semi-annual, two-business-day for large contributions), period start
    and end dates, when the candidate actually filed
  - declared_totals: the summary-page totals the candidate reported
    (total contributions, total expenditures, cash on hand start and end)
  - contributions: every individual contribution line with amount, date,
    contributor name and address, occupation and employer if stated,
    in-kind flag and description if applicable
  - expenditures: every expenditure line with payee, purpose, amount, date

RULES

1. Names: Copy exactly as written. Do not normalize spellings or
   capitalization. Identity resolution happens downstream.

2. filing_type classification: read the form header carefully.
   - "Pre-Election Report" / "Pre-Primary" / "Pre-Runoff" → pre_election
   - "Post-Election Report" / "Post-Primary" / "Post-Runoff" → post_election
   - "Semi-Annual Report" / "Mid-Year" / "Year-End" → semi_annual
   - "Two Business Day" / "$1,000+ Contributions" → two_business_day_large
   - "Amended" or "Amendment" → amendment
   - "Final Report" → final
   - Anything else → other

3. Dates: convert to YYYY-MM-DD. If the year is implied (e.g. "October
   15" on a 2024 cycle filing), infer from context. If a date is
   unreadable, leave it null — do not guess.

4. Amounts: parse as dollars. Strip "$" and commas. Negative numbers
   (refunds, returned contributions) keep their sign.

5. In-kind contributions: set is_in_kind=true and capture
   in_kind_description. Their amount is the dollar value as stated on
   the filing.

6. contributor_type:
   - "Individual" / personal name → individual
   - "Corporation" / "Inc" / "LLC" → corporation
   - "PAC" / "Political Action Committee" → pac
   - "Union" / labor organization → union
   - "Party Committee" → party
   - "Candidate's own funds" / "Loan from candidate" → self_funded
   - Anything else → other

7. Reconciliation: after extracting, internally check that sum of
   individual contributions matches declared_totals.total_contributions.
   If they disagree by more than $1, note the discrepancy in
   extraction_notes — do NOT alter the data.

8. If a field is unreadable or unclear, use null/None. Do not guess.

9. extraction_confidence:
   - "high" if every field was clearly readable and totals reconciled
   - "medium" if some fields were unclear or totals were slightly off
   - "low" if significant content was illegible or totals were way off
"""


OUTPUT_FORMAT_RULES = """\
Return exactly one JSON object matching the schema below. Do not include
any prose, markdown fences, or text outside the JSON. The response should
BEGIN with `{` and END with `}`. Use null for unknown/illegible fields.
Use empty arrays for empty lists.
"""


def _schema_for_prompt() -> str:
    import json as _json
    return _json.dumps(CampaignFilingExtraction.model_json_schema(), indent=2)


TEXT_INSTRUCTIONS = f"""\
You are extracting structured campaign finance data from a Campaign
Contribution Disclosure Report. The text below was produced from the
filing's text layer.

{_RULES}

{OUTPUT_FORMAT_RULES}

JSON SCHEMA:
{_schema_for_prompt()}

CCDR TEXT (page breaks marked):
"""


VISION_INSTRUCTIONS = f"""\
You are extracting structured campaign finance data from a scanned
Campaign Contribution Disclosure Report PDF.

Read the entire document carefully — especially handwritten amounts and
dates, which clerks often write by hand on photocopied state forms.

{_RULES}

{OUTPUT_FORMAT_RULES}

JSON SCHEMA:
{_schema_for_prompt()}

Now extract from the attached PDF.
"""


TEXT_MODEL = "claude-haiku-4-5"
VISION_MODEL = "claude-sonnet-4-6"


# =====================================================================
# Content-type dispatch (reuses the same shape as extractors/agendas.py)
# =====================================================================

PDF_CT = "application/pdf"
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_CT = "application/msword"


def _normalize_ct(content_type: str | None) -> str:
    return (content_type or "").lower().split(";", 1)[0].strip()


def _sniff_ct_from_magic(path: Path) -> str | None:
    head = path.read_bytes()[:8]
    if head.startswith(b"%PDF"):
        return PDF_CT
    if head.startswith(b"PK\x03\x04"):
        return DOCX_CT
    if head.startswith(b"\xd0\xcf\x11\xe0"):
        return DOC_CT
    return None


def extract_from_document(
    path: Path,
    content_type: str | None = None,
) -> tuple[CampaignFilingExtraction, str]:
    """Dispatch by content type. Returns (extraction, method).
    Method ∈ {'text_layer','vision','docx_text','doc_libreoffice'}.
    """
    ct = _normalize_ct(content_type)
    if ct not in (PDF_CT, DOCX_CT, DOC_CT):
        sniffed = _sniff_ct_from_magic(path)
        if sniffed:
            ct = sniffed

    if ct == PDF_CT:
        return _extract_from_pdf(path)
    if ct == DOCX_CT:
        return _extract_from_docx(path), "docx_text"
    if ct == DOC_CT:
        return _extract_from_doc(path), "doc_libreoffice"

    raise RuntimeError(
        f"unsupported content type {content_type!r} for {path.name} "
        f"(magic={path.read_bytes()[:4]!r})"
    )


def _extract_from_pdf(pdf_path: Path) -> tuple[CampaignFilingExtraction, str]:
    """Try tier-1 text layer, fall back to vision for scanned/handwritten forms."""
    from .pdf_text import _extract_text_layer, CONTENT_CHAR_THRESHOLD
    try:
        text, _pages, content_chars = _extract_text_layer(pdf_path)
        if content_chars >= CONTENT_CHAR_THRESHOLD:
            return _extract_from_text(text), "text_layer"
    except Exception:
        pass
    return _extract_from_pdf_vision(pdf_path), "vision"


def _extract_from_docx(doc_path: Path) -> CampaignFilingExtraction:
    from docx import Document
    doc = Document(doc_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((c.text or "").strip() for c in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return _extract_from_text("\n".join(parts))


def _extract_from_doc(doc_path: Path) -> CampaignFilingExtraction:
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if binary is None:
        raise RuntimeError(
            "DOC extraction requires libreoffice. "
            "Install on Mac: brew install --cask libreoffice. "
            "Install on Linux: apt-get install libreoffice-core libreoffice-writer."
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        subprocess.run(
            [binary, "--headless", "--convert-to", "txt:Text",
             "--outdir", tmpdir, str(doc_path)],
            check=True, timeout=180, capture_output=True,
        )
        txt_path = Path(tmpdir) / (doc_path.stem + ".txt")
        if not txt_path.exists():
            raise RuntimeError(f"libreoffice did not produce {txt_path.name}")
        text = txt_path.read_text(encoding="utf-8", errors="replace")
    return _extract_from_text(text)


def _parse_json_response(response) -> CampaignFilingExtraction:
    import json as _json
    import re as _re

    text = ""
    for block in response.content:
        if getattr(block, "type", "") == "text":
            text = block.text
            break
    fence = _re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, _re.DOTALL)
    if fence:
        text = fence.group(1)
    first = text.find("{")
    last = text.rfind("}")
    if first == -1 or last == -1 or last < first:
        raise ValueError(f"no JSON object in response: {text[:200]!r}")
    return CampaignFilingExtraction.model_validate(_json.loads(text[first : last + 1]))


def _extract_from_text(text: str) -> CampaignFilingExtraction:
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model=TEXT_MODEL,
        max_tokens=16384,
        messages=[{"role": "user", "content": TEXT_INSTRUCTIONS + "\n" + text}],
    )
    return _parse_json_response(response)


def _extract_from_pdf_vision(pdf_path: Path) -> CampaignFilingExtraction:
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    pdf_b64 = base64.standard_b64encode(pdf_path.read_bytes()).decode("utf-8")
    response = client.messages.create(
        model=VISION_MODEL,
        max_tokens=16384,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_b64,
                        },
                    },
                    {"type": "text", "text": VISION_INSTRUCTIONS},
                ],
            },
        ],
    )
    return _parse_json_response(response)
