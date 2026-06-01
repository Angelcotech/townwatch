"""
Extract structured docket items from meeting-agenda PDFs.

Agendas are the upstream half of the decision-making pipeline. They list
what WILL be heard — rezoning applications, variance requests, public
hearings, staff presentations — along with the applicant, recommended
action, and supporting documents. Minutes (extractors/minutes.py)
capture how the body acted; agendas capture what was asked.

Both extractors run automatically for every meeting that has the
corresponding PDF. When both succeed for the same meeting, a downstream
linker job (deferred) joins motions to their originating agenda items
so the full request → action trail is queryable.

Two-tier extraction matches minutes.py:
  1. pdfplumber → text layer (digital PDFs, free, instant)
  2. Claude vision on the raw PDF (scanned or fallback)

We deliberately do NOT use OCR (ocrmypdf/Tesseract) as a middle tier —
same reason as minutes.py: OCR errors cascade through identity
resolution. Vision reads original glyphs.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Literal

import anthropic
from pydantic import BaseModel, Field

from ..config import ANTHROPIC_API_KEY, VISION_RENDER_DPI
from .mistral_ocr import ocr_pdf
from .rasterize import vision_content
from .recovery import ExtractionReport, build_source, extract_with_ladder


# =====================================================================
# Pydantic schema — what Claude is constrained to return
# =====================================================================

ItemType = Literal[
    "rezoning",
    "variance",
    "special_exception",
    "conditional_use",
    "subdivision",
    "annexation",
    "ordinance_amendment",
    "public_hearing",
    "consent",
    "presentation",
    "old_business",
    "new_business",
    "staff_report",
    "procedural",
    "other",
]
HearingStatus = Literal[
    "first_reading",
    "second_reading",
    "public_hearing",
    "continued",
    "rescheduled",
    "not_applicable",
]
RecommendedAction = Literal[
    "approve",
    "deny",
    "table",
    "continue",
    "no_recommendation",
    "informational",
]
MeetingType = Literal["regular", "special", "workshop", "emergency", "executive_session"]
Confidence = Literal["high", "medium", "low"]


class AgendaMeetingMeta(BaseModel):
    date: str = Field(description="YYYY-MM-DD format")
    body_name: str = Field(description="The governing body (e.g. 'Planning Commission')")
    meeting_type: MeetingType
    scheduled_start_at: str | None = Field(default=None, description="HH:MM if printed on the agenda")
    extraction_confidence: Confidence


class AgendaItemRecord(BaseModel):
    item_number: str | None = Field(default=None, description="e.g. '7A', 'IV.2', or null if unnumbered")
    title: str = Field(description="Short docket title as written")
    description: str | None = Field(
        default=None,
        description="Full staff write-up or item paragraph if the agenda contains one. "
                    "Skip if the agenda only shows a one-line title."
    )
    item_type: ItemType
    applicant_name: str | None = Field(
        default=None,
        description="The person/entity who applied or petitioned for this item — developer, "
                    "business, resident, or staff (e.g. 'Acme Development LLC', 'John Smith, owner'). "
                    "For staff-initiated items (ordinances, presentations) this may be a department."
    )
    recommended_action: RecommendedAction | None = Field(
        default=None,
        description="What staff is asking the body to do for this item. Use 'informational' "
                    "for presentations / reports that don't require a vote. Null if not stated."
    )
    hearing_status: HearingStatus | None = Field(
        default=None,
        description="Procedural status if printed: first/second reading, public hearing, "
                    "continued/rescheduled item. Null when not applicable."
    )
    locations: list[str] = Field(
        default_factory=list,
        description="Every property location mentioned, one entry per distinct location. "
                    "Use the original label as written — include the full string with "
                    "address, parcel ID, and any descriptive prefix "
                    '(e.g. "1110 Dodge Lane (Parcel ID 070 009)").',
    )
    documents_referenced: list[str] = Field(
        default_factory=list,
        description="Attachments, exhibits, staff reports, or supporting documents listed for this item.",
    )
    source_page: int = Field(description="1-indexed PDF page where this item begins")


class AgendaExtraction(BaseModel):
    meeting: AgendaMeetingMeta
    agenda_items: list[AgendaItemRecord]
    document_summary: str = Field(
        default="",
        description=(
            "2-3 sentence plain-English summary of what this agenda contains overall. "
            "Written for a citizen who has not opened the PDF. Focus on the substantive "
            "items the body will hear (rezoning, variance, contracts, budget) and skip "
            "procedural boilerplate. Example: 'The Planning Commission will hear three "
            "rezoning applications, including a downtown redevelopment, plus a routine "
            "variance for a backyard setback. Two items are returning from prior meetings.'"
        ),
    )
    extraction_notes: str = Field(
        default="",
        description="Ambiguities, illegible passages, or unusual items. Empty string if clean.",
    )


# =====================================================================
# Prompts
# =====================================================================

_RULES = """\
WHAT TO EXTRACT
- Every substantive docket item: rezoning, variance, special exception, conditional use, subdivision, annexation, ordinance amendments, public hearings, contracts up for action, presentations
- Applicant identity for every item (who petitioned/applied — name + entity if both stated)
- Recommended action from staff (approve / deny / table / continue / informational)
- Hearing status when printed (first reading, second reading, public hearing, continued, rescheduled)
- Every property address, parcel ID, or described area mentioned per item
- Documents/exhibits/staff reports referenced per item
- Source page (1-indexed) where each item begins

WHAT TO SKIP
- Call to order, pledge of allegiance, roll call (just the procedural header)
- "Approval of minutes" line items
- Adjournment
- Generic "public comment" / "citizen input" header lines unless tied to a specific item
- Watermarks, page numbers, draft labels

ITEM TYPING — use these signals
- "Rezoning" / "Zoning Map Amendment" / "ZMA" / "Map Amendment" → rezoning
- "Variance" → variance
- "Special Exception" / "Special Use Permit" → special_exception
- "Conditional Use Permit" / "CUP" → conditional_use
- "Subdivision" / "Plat" → subdivision
- "Annexation" → annexation
- "Ordinance" amending text of the code → ordinance_amendment
- "Public Hearing" without a more specific category → public_hearing
- "Consent Agenda" items → consent (extract each one)
- Reports/updates/presentations not requiring action → presentation
- Carried-over items not falling above → old_business
- Brand-new items not falling above → new_business
- Internal department updates → staff_report
- Procedural votes (approve agenda, adjourn) → procedural
- Anything else → other

DOCUMENT SUMMARY
- Always populate document_summary with a 2-3 sentence plain-English overview of the agenda. This is what a citizen will read instead of opening the PDF. Write it AFTER you have surveyed the items so it accurately reflects what's on the docket.
- Focus on substantive items (rezoning, variance, contracts, budget). Skip procedural boilerplate (roll call, approval of minutes, adjournment).
- Be concrete: name the type of item, applicants when notable, and any flagged returning items.

RULES
1. Names: Copy exactly as written. Do not normalize spellings or capitalization. Identity resolution happens downstream.
2. applicant_name: who REQUESTED or APPLIED for this item. Common signals: "Applicant:", "Petitioner:", "On the application of...", "Request from...", "Filed by...". For internal items (ordinances, code amendments) the applicant may be a department or staff member.
3. recommended_action: only set when staff recommendation is explicit. "Staff recommends approval" → approve. "Recommended for denial" → deny. Information-only items → informational. Null when the agenda doesn't state a recommendation.
4. hearing_status: only set when the agenda labels it. Don't infer from item type.
5. locations: every distinct address / parcel / described area per item. One entry per location, full original label.
6. source_page: PDF page (1-indexed) where the item starts.
7. If a field is illegible or unclear, use null/None. Do not guess.
8. extraction_confidence: "high" if every item was clearly readable, "medium" if some details were unclear, "low" if significant content was illegible.

ATTRIBUTION TO DECISION — IMPORTANT
The agenda lists what is BEING PROPOSED, not what was decided. Do NOT
guess outcomes. Outcome data lives in the minutes (separate extractor).
"""


OUTPUT_FORMAT_RULES = """\
Return exactly one JSON object matching the schema below. Do not include any
prose, markdown fences, or text outside the JSON. The response should BEGIN
with `{` and END with `}`. Use null for unknown/illegible fields. Use empty
arrays for empty lists.
"""


def _schema_for_prompt() -> str:
    """Compact JSON-schema description for the prompt."""
    import json as _json
    return _json.dumps(AgendaExtraction.model_json_schema(), indent=2)


TEXT_INSTRUCTIONS = f"""\
You are extracting structured docket data from a meeting agenda. The text below was produced from the PDF's text layer.

{_RULES}

{OUTPUT_FORMAT_RULES}

JSON SCHEMA:
{_schema_for_prompt()}

MEETING AGENDA TEXT (page breaks marked):
"""

VISION_INSTRUCTIONS = f"""\
You are extracting structured docket data from a scanned meeting agenda PDF.

Read the entire document carefully.

{_RULES}

{OUTPUT_FORMAT_RULES}

JSON SCHEMA:
{_schema_for_prompt()}

Now extract from the attached PDF.
"""

# Models match minutes.py: Haiku for text layer, Sonnet vision for everything else.
TEXT_MODEL = "claude-haiku-4-5"
VISION_MODEL = "claude-sonnet-4-6"

# CivicEngage sometimes serves placeholder PDFs for older meetings whose
# agendas were never actually uploaded — valid PDF header, zero-content
# pages, ~1.5-2KB total. Calling vision on these wastes Sonnet calls only
# to get items=[]. If a PDF is below this size AND has no text layer, we
# short-circuit and return an empty extraction marked as a stub.
STUB_PDF_SIZE_BYTES = 5_000


# =====================================================================
# API calls
# =====================================================================

def extract_from_pdf(pdf_path: Path) -> tuple[AgendaExtraction, str, ExtractionReport]:
    """
    Returns (extraction, method, report) where method is 'text_layer',
    'vision', or 'stub_skipped'.

    Tiny PDFs with no text layer are recognised as CivicEngage placeholder
    stubs and skipped without an API call. Everything else runs through the
    escalating-filter ladder (recovery.extract_with_ladder): page-windowed,
    with primary → retry → sub-chunk → cross-strategy → pdf-repair recovery
    and classified anomalies in the report. No OCR tier — OCR errors break
    identity resolution.
    """
    text_layer = extract_text_layer_only(pdf_path)
    # Stub detection BEFORE OCR: empty text layer + tiny file = placeholder
    # PDF, return an empty extraction with no API call.
    if text_layer is None and pdf_path.stat().st_size <= STUB_PDF_SIZE_BYTES:
        return _stub_extraction(pdf_path), "stub_skipped", ExtractionReport(total_units=1, clean=1, method="stub_skipped")

    method = "text_layer"
    if text_layer is None:
        # Scanned: Mistral OCR primary; vision fallback if OCR yields nothing.
        text_layer = ocr_pdf(pdf_path)
        method = "ocr" if text_layer else "vision"
    source = build_source(pdf_path, text_layer)
    extraction, report = extract_with_ladder(
        source,
        text_window_fn=_extract_text_window,
        vision_window_fn=_extract_vision_window,
        merge_fn=_merge_extractions,
    )
    report.method = method
    return extraction, method, report


def _stub_extraction(pdf_path: Path) -> AgendaExtraction:
    """Return an empty extraction without an API call.

    Used when the PDF is a CivicEngage placeholder stub (tiny file, no
    text layer). Callers can detect via extraction_notes or method.
    """
    return AgendaExtraction(
        meeting=AgendaMeetingMeta(
            date="1900-01-01",  # caller already knows the real date
            body_name="",
            meeting_type="regular",
            extraction_confidence="low",
        ),
        agenda_items=[],
        extraction_notes=(
            f"stub_skipped: PDF size {pdf_path.stat().st_size} bytes has no text layer; "
            f"recognised as a CivicEngage placeholder stub (no real content uploaded)."
        ),
    )


# =====================================================================
# Content-type dispatch — agendas arrive as PDF / DOCX / DOC depending
# on jurisdiction vintage. This is the entry point most callers should
# use; extract_from_pdf is preserved for backward compatibility.
# =====================================================================

DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_CT = "application/msword"
PDF_CT = "application/pdf"


def _normalize_ct(content_type: str | None) -> str:
    return (content_type or "").lower().split(";", 1)[0].strip()


def _sniff_ct_from_magic(path: Path) -> str | None:
    """Sniff content type from the first bytes. Used when HTTP didn't tell us."""
    head = path.read_bytes()[:8]
    if head.startswith(b"%PDF"):
        return PDF_CT
    if head.startswith(b"PK\x03\x04"):  # ZIP container — likely DOCX
        return DOCX_CT
    if head.startswith(b"\xd0\xcf\x11\xe0"):  # OLE Compound Document — DOC, XLS, PPT
        return DOC_CT
    return None


def extract_from_document(
    path: Path,
    content_type: str | None = None,
) -> tuple[AgendaExtraction, str, ExtractionReport]:
    """Dispatch extraction by content type.

    Returns (extraction, method, report) where method is one of
    'text_layer', 'vision', 'stub_skipped', 'docx_text', 'doc_libreoffice'.
    DOCX/DOC/stub paths return a trivially-clean report; the PDF path returns
    the recovery ladder's report.

    content_type may be passed straight from an HTTP Content-Type header;
    if absent or unrecognised, we sniff the file's magic bytes. Raises
    RuntimeError for unsupported formats so the caller can mark the
    meeting unavailable cleanly.
    """
    ct = _normalize_ct(content_type)
    if ct not in (PDF_CT, DOCX_CT, DOC_CT):
        sniffed = _sniff_ct_from_magic(path)
        if sniffed:
            ct = sniffed

    if ct == PDF_CT:
        return extract_from_pdf(path)
    if ct == DOCX_CT:
        return _extract_from_docx(path), "docx_text", ExtractionReport(total_units=1, clean=1)
    if ct == DOC_CT:
        return _extract_from_doc(path), "doc_libreoffice", ExtractionReport(total_units=1, clean=1)

    raise RuntimeError(
        f"unsupported document type {content_type!r} for {path.name} "
        f"(magic={path.read_bytes()[:4]!r})"
    )


def _extract_from_docx(doc_path: Path) -> AgendaExtraction:
    """Read DOCX paragraphs + tables as plain text, then feed to Haiku.

    DOCX is structured XML — no vision needed. Tables matter because
    many older agenda templates put the docket in a table rather than
    paragraph form.
    """
    from docx import Document  # lazy import — python-docx
    doc = Document(doc_path)
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join((c.text or "").strip() for c in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    text = "\n".join(parts)
    return _extract_text_window(text)


def _extract_from_doc(doc_path: Path) -> AgendaExtraction:
    """Shell to libreoffice headless to convert legacy DOC → text, then Haiku.

    DOC is the binary Microsoft Word format; no good native Python
    reader exists. libreoffice is the standard converter on both Mac
    (brew install --cask libreoffice) and Linux (apt install libreoffice).
    """
    binary = shutil.which("libreoffice") or shutil.which("soffice")
    if binary is None:
        raise RuntimeError(
            "DOC extraction requires libreoffice. "
            "Install on Mac: brew install --cask libreoffice. "
            "Install on Linux: apt-get install libreoffice-core libreoffice-writer."
        )
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run(
                [binary, "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(doc_path)],
                check=True, timeout=180, capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(
                f"libreoffice conversion failed: {e.stderr.decode('utf-8', errors='replace')[:500]}"
            ) from e
        txt_path = Path(tmpdir) / (doc_path.stem + ".txt")
        if not txt_path.exists():
            raise RuntimeError(f"libreoffice did not produce {txt_path.name}")
        text = txt_path.read_text(encoding="utf-8", errors="replace")
    return _extract_text_window(text)


def extract_text_layer_only(pdf_path: Path) -> str | None:
    """Tier 1 only: pdfplumber on the existing text layer. None if no real text."""
    from .pdf_text import _extract_text_layer, CONTENT_CHAR_THRESHOLD
    try:
        text, _pages, content_chars = _extract_text_layer(pdf_path)
    except Exception:
        return None
    if content_chars < CONTENT_CHAR_THRESHOLD:
        return None
    return text


def _parse_json_response(response) -> AgendaExtraction:
    """Locate and validate the JSON object in Claude's response."""
    import json as _json
    import re as _re

    text = ""
    for block in response.content:
        if getattr(block, "type", "") == "text":
            text = block.text
            break

    fence_match = _re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, _re.DOTALL)
    if fence_match:
        text = fence_match.group(1)

    first = text.find("{")
    last = text.rfind("}")
    if first == -1 or last == -1 or last < first:
        stop = getattr(response, "stop_reason", "?")
        raise ValueError(f"No JSON object found in response (stop_reason={stop}): {text[:200]!r}")
    return AgendaExtraction.model_validate(_json.loads(text[first : last + 1]))


def _extract_text_window(text: str) -> AgendaExtraction:
    """Extract one window of agenda text (Haiku → JSON), streamed with a 32K
    budget. The page-window ladder maps this over large documents."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    with client.messages.stream(
        model=TEXT_MODEL,
        max_tokens=32000,
        messages=[
            {"role": "user", "content": TEXT_INSTRUCTIONS + "\n" + text},
        ],
    ) as stream:
        response = stream.get_final_message()
    return _parse_json_response(response)


def _extract_vision_window(pdf_path: Path, dpi: int | None = VISION_RENDER_DPI) -> AgendaExtraction:
    """Extract one window sub-PDF via Sonnet vision, streamed with a 32K budget
    (thinking + output share it — streaming lifts the SDK's 10-min non-stream
    guard). The ladder maps this over large documents. ``dpi`` controls
    rasterization: None ships the raw PDF (default), an int rasterizes pages to
    images at that resolution. Defaults to config.VISION_RENDER_DPI."""
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    with client.messages.stream(
        model=VISION_MODEL,
        max_tokens=32000,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        messages=[{"role": "user", "content": vision_content(pdf_path, VISION_INSTRUCTIONS, dpi=dpi)}],
    ) as stream:
        response = stream.get_final_message()
    return _parse_json_response(response)


def _merge_extractions(partials: list[tuple[AgendaExtraction, int]]) -> AgendaExtraction:
    """Reduce per-window agenda extractions into one. Items are offset-
    corrected to document page numbers and de-duplicated by (page, title);
    meeting metadata comes from the first window; the document summary is
    synthesized from the per-window summaries."""
    first = partials[0][0]
    items: list[AgendaItemRecord] = []
    seen: set[tuple[int, str]] = set()
    notes: list[str] = []
    summaries: list[str] = []
    for ext, offset in partials:
        for it in ext.agenda_items:
            it.source_page = it.source_page + offset
            key = (it.source_page, (it.title or "").strip().casefold())
            if key in seen:
                continue
            seen.add(key)
            items.append(it)
        if ext.extraction_notes.strip():
            notes.append(ext.extraction_notes.strip())
        if ext.document_summary.strip():
            summaries.append(ext.document_summary.strip())
    items.sort(key=lambda it: it.source_page)
    return AgendaExtraction(
        meeting=first.meeting,
        agenda_items=items,
        document_summary=_synthesize_summary(summaries),
        extraction_notes=" | ".join(notes)[:2000],
    )


def _synthesize_summary(summaries: list[str]) -> str:
    """Collapse per-window agenda summaries into one 2-3 sentence summary
    (one cheap Haiku call). Falls back to the longest window summary on
    failure so a reduce hiccup never loses the summary."""
    summaries = [s for s in summaries if s.strip()]
    if not summaries:
        return ""
    if len(summaries) == 1:
        return summaries[0]
    joined = "\n".join(f"- {s}" for s in summaries)
    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        resp = client.messages.create(
            model=TEXT_MODEL,
            max_tokens=512,
            messages=[{"role": "user", "content": (
                "Below are per-section summaries of one government meeting agenda. "
                "Write a single 2-3 sentence plain-English summary for a citizen of "
                "what's on the docket. Return only the summary.\n\n" + joined
            )}],
        )
        for block in resp.content:
            if getattr(block, "type", "") == "text" and block.text.strip():
                return block.text.strip()
    except Exception:
        pass
    return max(summaries, key=len)
